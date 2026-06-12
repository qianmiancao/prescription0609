# --- 1. 核心兼容性补丁 (必须处于文件最顶部，且在任何 Langchain 导入之前) ---
import sys

try:
    import pysqlite3
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    # 如果是本地 Windows 环境没有 pysqlite3-binary，则跳过
    pass

import streamlit as st
import os
import json
import warnings
import tempfile
import pandas as pd
import re 
import shutil
from datetime import datetime
from io import BytesIO
from docx import Document 
from docx.shared import Pt

# --- 2. 基础环境配置 ---
warnings.filterwarnings("ignore", category=DeprecationWarning)
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader

# 数据库存储路径 (Streamlit Cloud 环境建议使用新路径刷新)
DB_PATH = "./drug_db_v_final" 

# --- 3. 逻辑类定义 ---

class KnowledgeManager:
    def __init__(self, model_name, db_path):
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name, model_kwargs={'device': 'cpu'})
        self.db_path = db_path
        self.vectorstore = self.init_db()

    def init_db(self):
        """安全初始化数据库，若报错则强制重置"""
        try:
            return Chroma(persist_directory=self.db_path, embedding_function=self.embeddings)
        except Exception:
            if os.path.exists(self.db_path):
                shutil.rmtree(self.db_path)
            return Chroma(persist_directory=self.db_path, embedding_function=self.embeddings)

    def upload_docs(self, file_path, original_name):
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        for doc in docs:
            doc.metadata = {"source": original_name} 
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        self.vectorstore.add_documents(splitter.split_documents(docs))
        return len(docs)

    def retrieve_context(self, drug_names):
        results = []
        sources = set()
        for name in drug_names:
            if not name: continue
            # 搜索最相关的片段
            docs = self.vectorstore.similarity_search(name, k=3)
            for d in docs:
                src = os.path.basename(d.metadata.get('source', '资料'))
                # 关键词简单校验，防止跨药匹配
                if name[:2] in d.page_content or name[:2] in src:
                    results.append(f"【内容参考《{src}》】: {d.page_content}")
                    sources.add(src)
        return "\n\n".join(results), list(sources)

class PharmacyAgent:
    def __init__(self, api_key):
        self.llm = ChatOpenAI(
            model="deepseek-chat", 
            api_key=api_key, 
            base_url="https://api.deepseek.com", 
            temperature=0.1
        )

    def audit(self, p_data, context):
        # 强制要求纯文本格式，避免表格和星号
        system_prompt = """你是一位资深临床药师。请根据参考资料审核处方。
        
        【格式规范】：
        1. 严禁使用 * 符号进行加粗。
        2. 严禁使用 | 或 - 等符号构建表格，请直接使用纯文本列表。
        3. “产品明细”请按此格式列出：
           1. 药品：[名称]；剂量：[用量]；频次：[频次]；用法：[用法]
        4. 点评结论分为：一、基本信息；二、产品明细；三、适宜性点评；四、药师结论。
        
        【审核逻辑】：
        - 若是儿童，需依据体重计算 mg/kg 剂量。
        - 若是成人，体重为空时，请依据成人常规剂量审核。
        """
        
        prompt = ChatPromptTemplate.from_template("指令: {sys}\n参考资料: {ctx}\n处方数据: {rx}")
        chain = prompt | self.llm
        res = chain.invoke({
            "sys": system_prompt,
            "ctx": context,
            "rx": json.dumps(p_data, ensure_ascii=False)
        })
        # 二次清洗残留星号和表格线
        return res.content.replace('*', '').replace('|', '').strip()

# --- 4. 辅助函数 ---

def generate_docx(text, p_no):
    doc = Document()
    doc.add_heading('临床处方点评报告', 0)
    doc.add_paragraph(f"处方编号：{p_no}")
    doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("-" * 30)
    for line in text.split('\n'):
        if line.strip():
            para = doc.add_paragraph(line.strip())
            if any(line.startswith(p) for p in ["一、", "二、", "三、", "四、"]):
                para.runs[0].bold = True
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 5. Streamlit UI 界面 ---

def main():
    st.set_page_config(page_title="AI 药师点评专家系统", layout="wide", page_icon="💊")
    
    # 使用缓存避免重复加载嵌入模型
    @st.cache_resource(show_spinner="药学大脑启动中...")
    def get_km(v="1.0"):
        return KnowledgeManager("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2", DB_PATH)
    
    km = get_km()

    if "rpt" not in st.session_state: st.session_state.rpt = ""
    if "srcs" not in st.session_state: st.session_state.srcs = []

    with st.sidebar:
        st.header("⚙️ 知识库管理")
        api_key = st.text_input("DeepSeek API Key:", type="password")
        files = st.file_uploader("上传药品说明书 (PDF)", accept_multiple_files=True)
        if files and st.button("✨ 同步到库"):
            with st.spinner("建立索引中..."):
                for f in files:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(f.getvalue())
                        km.upload_docs(tmp.name, f.name)
            st.success("同步成功")
        
        if st.button("🗑️ 清空所有记录"):
            if os.path.exists(DB_PATH): shutil.rmtree(DB_PATH)
            st.rerun()

    st.title("🏥 临床药师点评专家平台")

    col1, col2 = st.columns([1, 1.2])

    with col1:
        st.subheader("📋 处方录入")
        with st.container(border=True):
            p_no = st.text_input("处方流水号", value="RX"+datetime.now().strftime("%H%M%S"))
            diag = st.text_input("临床诊断", "急性支气管炎")
            
            p_type = st.radio("患者类型", ["成人", "儿童"], horizontal=True)
            
            c1, c2 = st.columns(2)
            age = c1.number_input("年龄", value=35 if p_type == "成人" else 6)
            
            # 成人无需输入体重逻辑
            weight = None
            if p_type == "儿童":
                weight = c2.number_input("体重 (kg)", value=22.0)
            else:
                c2.info("💡 成人模式免填体重")

            st.markdown("**药品清单明细**")
            med_df = st.data_editor(
                pd.DataFrame([{"产品名称": "阿莫西林胶囊", "剂量": "0.5g", "频次": "TID", "用法": "口服"}]),
                num_rows="dynamic", use_container_width=True
            )

        if st.button("🔍 执行深度点评", type="primary", use_container_width=True):
            if not api_key:
                st.error("请输入 API Key")
            else:
                agent = PharmacyAgent(api_key)
                with st.spinner("正在检索并进行逻辑推理..."):
                    drugs = med_df["产品名称"].tolist()
                    ctx, srcs = km.retrieve_context(drugs)
                    st.session_state.srcs = srcs
                    
                    p_data = {
                        "no": p_no, "diag": diag, "type": p_type,
                        "age": age, "weight": weight, "meds": med_df.to_dict('records')
                    }
                    st.session_state.rpt = agent.audit(p_data, ctx)
                    st.rerun()

    with col2:
        st.subheader("📝 审核报告 (实时同步)")
        if st.session_state.srcs:
            st.success(f"参考溯源: {', '.join(st.session_state.srcs)}")
        
        if st.session_state.rpt:
            # 实时更新内容
            st.session_state.rpt = st.text_area("报告内容:", value=st.session_state.rpt, height=550)
            
            st.divider()
            bt1, bt2 = st.columns(2)
            with bt1:
                docx_data = generate_docx(st.session_state.rpt, p_no)
                st.download_button(
                    "📥 导出 Word 报告", 
                    data=docx_data, 
                    file_name=f"点评报告_{p_no}.docx",
                    use_container_width=True,
                    type="primary"
                )
            with bt2:
                if st.button("🗑️ 清空内容", use_container_width=True):
                    st.session_state.rpt = ""
                    st.rerun()
        else:
            st.info("👈 请在左侧录入处方并执行点评。")

if __name__ == "__main__":
    main()
